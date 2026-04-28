"""4479 マクアケ Deep Research V2（決算説明資料完全反映版） MD → DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH   = "outputs/fukai/4479_マクアケ_WATCH_DeepResearch_20260428_v2.md"
DOCX_PATH = "outputs/fukai/4479_マクアケ_WATCH_DeepResearch_20260428_v2.docx"

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_NAVY  = (0,   70,  127)
C_BLUE  = (31,  73,  125)
C_DEEP  = (54,  96,  146)
C_WHITE = (255, 255, 255)
C_GREEN = (0,   112,  0)
C_RED   = (192,  0,   0)


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "Yu Gothic")
    rPr.insert(0, rf)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: C_NAVY, 2: C_BLUE, 3: C_DEEP}
    sizes  = {1: 16, 2: 13, 3: 11.5}
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=colors.get(level, (0, 0, 0)))
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if level <= 2:
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "4472C4")
        pBdr.append(bot)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=10, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)


def add_table_from_md(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = t.cell(i, j)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            cell.text = clean
            runs = cell.paragraphs[0].runs
            if runs:
                set_font(runs[0], size=9, bold=(i == 0))
                if i == 0:
                    runs[0].font.color.rgb = RGBColor(*C_WHITE)
            if i == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F497D")
                tcPr.append(shd)
            elif i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DEEAF1")
                tcPr.append(shd)
    doc.add_paragraph()


# ── カバーページ ──
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("オーキッドツリーキャピタル アナリストチーム")
set_font(r, size=11, color=C_BLUE)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("4479 株式会社マクアケ")
set_font(r, size=20, bold=True, color=C_NAVY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Deep Research — BUY +72  V2（決算説明資料完全反映・中計前倒し達成確定版）")
set_font(r, size=14, bold=True, color=C_BLUE)

doc.add_paragraph()
doc.add_paragraph()

cover_items = [
    ("担当アナリスト", "深井"),
    ("分析日", "2026-04-28（2Q決算・説明資料開示当日・V2確定版）"),
    ("V2改訂理由", "決算説明資料解析。Q2 QoQ▲10.2%→季節性確定 / H2残りOP¥1.02億→¥3-4億先行投資と判明 / 中計前倒し達成→新中計発表が次の大カタリスト"),
    ("投資判断", "WATCH +58 → BUY +72（格上げ）打診25%即時可"),
    ("現在株価", "¥920（2Q決算開示当日）"),
    ("推奨アクション", "打診25%即時（¥920前後）。本玉は2026/9期末の新中計発表時"),
    ("Base TP（12ヶ月）", "¥1,300（+41.3%）← V1 ¥1,200から上方修正"),
    ("Bull TP", "¥1,800（+95.7%）— 新中計大幅上方改訂×グローバル展開本格化シナリオ"),
    ("Bear TP", "¥650（▲29.3%）— リピート率低下加速×競合圧迫×先行投資空振りシナリオ"),
    ("加重平均TP", "¥1,159（確率加重：Base50%+Bull30%+Bear20%）→ 現値比+26.0%"),
    ("V2 3弾丸（追加）", "①中計2期前倒し達成→新中計（2026/11）が株価再評価トリガー / ②GMV+22.5% YoY・テイクレート29.1%（高水準） / ③H2¥3-4億先行投資=来期OP急拡大のファンド"),
    ("Q2 GMV", "¥5,107M（+22.5% YoY・▲10.2% QoQ）→ 季節性確定"),
    ("Q2テイクレート", "29.1%（Q1 28.3%比+0.8pt改善）"),
    ("Q2 UU数", "+28.4% YoY（会員数+12.1%）"),
    ("通期修正後売上", "¥5,400M（修正前¥4,763M比+13.4%）"),
    ("通期修正後OP", "¥670〜800M（修正前¥400M比+67.5〜100%）"),
    ("V2最大リスク", "リピート率低下（サポーター70.7%▲2.6pt YoY・実行者59.3%▲2.4pt YoY）。先送り説明の実現性が2027/9期の最大検証項目"),
    ("FY2026/9予想EPS", "¥49〜58（通期OP¥670〜800Mベース）"),
    ("ForwardPER（Base）", "22.4x（¥1,300 / ¥58）→ 成長株として許容範囲"),
    ("中計達成状況", "FY2026/9通期見通し：売上¥5,400M・OP¥670〜800M / 中計目標：売上¥5,200M・OP¥700M → 売上は超過確定、OP中央値も達成圏"),
    ("新中計発表予定", "2026/9期末（2026年11月頃）— 次の大カタリスト"),
    ("損切り水準", "¥750（25MA大幅割れ、▲18.5%）"),
    ("IRコンタクト", "https://www.makuake.co.jp/ir/"),
]

for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}：{value}")
    set_font(r, size=10.5)

doc.add_page_break()

# ── 本文パース ──
i = 0
table_buf = []
in_code   = False
code_buf  = []

while i < len(lines):
    raw      = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_font(run, size=9)
            run.font.name = "Courier New"
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    if stripped.startswith("|"):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    if stripped.startswith("### "):
        add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        add_heading(stripped[2:], level=1)
    elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        add_body(stripped.strip("*"), bold=True)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        content = stripped[2:]
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=10, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10)
    elif stripped.startswith("> "):
        add_body(stripped[2:], indent=True)
    elif stripped == "" or stripped == "---":
        doc.add_paragraph()
    else:
        add_body(stripped)

    i += 1

if table_buf:
    add_table_from_md(table_buf)

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
