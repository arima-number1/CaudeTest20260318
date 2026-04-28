"""4620 藤倉化成 Deep Research MD → DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH   = "outputs/fukai/4620_藤倉化成_WATCH_DeepResearch_20260428.md"
DOCX_PATH = "outputs/fukai/4620_藤倉化成_WATCH_DeepResearch_20260428.docx"

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_NAVY  = (0,   70,  127)
C_BLUE  = (31,  73,  125)
C_DEEP  = (54,  96,  146)
C_WHITE = (255, 255, 255)
C_GREEN = (0,   112,  0)
C_RED   = (192,  0,   0)


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "Yu Gothic")
    rPr.insert(0, rf)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: C_NAVY, 2: C_BLUE, 3: C_DEEP}
    sizes  = {1: 16, 2: 13, 3: 11.5}
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=colors.get(level, (0, 0, 0)))
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if level <= 2:
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "4472C4")
        pBdr.append(bot)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=10, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)


def add_table_from_md(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = t.cell(i, j)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            cell.text = clean
            runs = cell.paragraphs[0].runs
            if runs:
                set_font(runs[0], size=9, bold=(i == 0))
                if i == 0:
                    runs[0].font.color.rgb = RGBColor(*C_WHITE)
            if i == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F497D")
                tcPr.append(shd)
            elif i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DEEAF1")
                tcPr.append(shd)
    doc.add_paragraph()


# ── カバーページ ──
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("オーキッドツリーキャピタル アナリストチーム")
set_font(r, size=11, color=C_BLUE)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("4620 藤倉化成株式会社")
set_font(r, size=20, bold=True, color=C_NAVY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Deep Research — WATCH +65（上方修正・増配反映版）ナフサリスク精査")
set_font(r, size=14, bold=True, color=C_BLUE)

doc.add_paragraph()
doc.add_paragraph()

cover_items = [
    ("担当アナリスト", "深井"),
    ("分析日", "2026-04-28（上方修正・増配発表当日）"),
    ("投資判断", "WATCH +65 / 打診20%可"),
    ("現在株価", "¥1,222（本日+5.89%、修正発表日）"),
    ("Base TP（12ヶ月）", "¥1,600（+31.0%）"),
    ("Bull TP", "¥2,000（+63.7%）— 持合株解消完了×ROE改善シナリオ"),
    ("Bear TP", "¥900（▲26.4%）— ナフサ調達危機×建設需要急落"),
    ("加重平均TP", "¥1,200（ほぼ現値。エッジはBase確率向上情報次第）"),
    ("FY2026/3 修正NI", "¥30億（前回予想¥23億比+30.4%）"),
    ("FY2026/3 修正EPS", "¥102.27（一時益含む）/ 正常化EPS推計¥61-68"),
    ("FY2026/3 配当", "¥20確定（中間¥9+期末¥11、+¥2増配）"),
    ("安全域", "現預金¥148.9億（1株¥512 = 株価の42%）EV/株¥710"),
    ("ナフサリスク判定", "現在〜6月: 大丈夫。7月以降: 要注意（旭化成「臨界点」と同期）"),
    ("PBR / BPS", "0.80x / ¥1,532（解散価値以下で割安）"),
    ("正常化EV/PE", "¥710(EV/株) ÷ ¥70(正常化EPS) = 10.1x"),
    ("損切り水準", "¥950（現値▲22.3%、現預金フロア¥512考慮）"),
    ("推奨アクション", "打診20%（¥1,200前後）。本玉は8月Q1決算でナフサ転嫁確認後"),
    ("IRコンタクト", "http://www.fkkasei.co.jp/ir/index.html"),
]

for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}：{value}")
    set_font(r, size=10.5)

doc.add_page_break()

# ── 本文パース ──
i = 0
table_buf = []
in_code   = False
code_buf  = []

while i < len(lines):
    raw      = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_font(run, size=9)
            run.font.name = "Courier New"
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    if stripped.startswith("|"):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    if stripped.startswith("### "):
        add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        add_heading(stripped[2:], level=1)
    elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        add_body(stripped.strip("*"), bold=True)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        content = stripped[2:]
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=10, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10)
    elif stripped.startswith("> "):
        add_body(stripped[2:], indent=True)
    elif stripped == "" or stripped == "---":
        doc.add_paragraph()
    else:
        add_body(stripped)

    i += 1

if table_buf:
    add_table_from_md(table_buf)

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
