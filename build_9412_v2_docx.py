"""9412 スカパーJSAT Deep Research V2（本決算実績反映版） MD → DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH   = "outputs/fukai/9412_スカパーJSAT_BUY_DeepResearch_20260428_v2.md"
DOCX_PATH = "outputs/fukai/9412_スカパーJSAT_BUY_DeepResearch_20260428_v2.docx"

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_NAVY  = (0,   70,  127)
C_BLUE  = (31,  73,  125)
C_DEEP  = (54,  96,  146)
C_WHITE = (255, 255, 255)
C_GREEN = (0,   112,  0)
C_RED   = (192,  0,   0)


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "Yu Gothic")
    rPr.insert(0, rf)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: C_NAVY, 2: C_BLUE, 3: C_DEEP}
    sizes  = {1: 16, 2: 13, 3: 11.5}
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=colors.get(level, (0, 0, 0)))
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if level <= 2:
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "4472C4")
        pBdr.append(bot)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=10, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)


def add_table_from_md(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = t.cell(i, j)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            clean = re.sub(r"✅|🚨|—", lambda m: m.group(), clean)
            cell.text = clean
            runs = cell.paragraphs[0].runs
            if runs:
                set_font(runs[0], size=9, bold=(i == 0))
                if i == 0:
                    runs[0].font.color.rgb = RGBColor(*C_WHITE)
            if i == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F497D")
                tcPr.append(shd)
            elif i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DEEAF1")
                tcPr.append(shd)
    doc.add_paragraph()


# ── カバーページ ──
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("オーキッドツリーキャピタル アナリストチーム")
set_font(r, size=11, color=C_BLUE)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("9412 スカパーJSATホールディングス")
set_font(r, size=20, bold=True, color=C_NAVY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Deep Research — BUY +90  V2（FY2026/3本決算実績反映・ガイダンス¥270億確認）")
set_font(r, size=14, bold=True, color=C_BLUE)

doc.add_paragraph()
doc.add_paragraph()

cover_items = [
    ("担当アナリスト", "深井"),
    ("分析日", "2026-04-28（本決算発表当日・V2確定版）"),
    ("V2改訂理由", "FY2026/3 NI¥233億確定（推計比+6-8%上振れ）+ FY2027/3 NI¥270億ガイダンス（コンセンサス¥198億比+36.4%）"),
    ("投資判断", "WATCH +72 → BUY +90（格上げ）"),
    ("現在株価", "¥3,195（本日+3.40%、決算日ポジティブプリント確認）"),
    ("FY2026/3 NI実績", "¥233億（+22% YoY）— 推計¥215-220億を上振れ"),
    ("FY2027/3 NI ガイダンス", "¥270億（+15.9%）— コンセンサス¥198億比+36.4%超過"),
    ("竹村トリガー判定", "✅ ¥270億 ≥ ¥210億 → 本玉60%追加断定（+28.6%余裕）"),
    ("推奨アクション", "本玉60%追加（計90%）。¥3,195前後・¥3,400まで追跡"),
    ("Base TP（12ヶ月）", "¥4,800（+50.2%）← V1 ¥4,100から+¥700上方修正"),
    ("Bull TP", "¥6,000（+87.8%）— 防衛省コンステレーション受注開示シナリオ"),
    ("Bear TP", "¥2,300（▲28.0%）— メディア崩壊×防衛失望（V1から変更なし）"),
    ("加重平均TP", "¥4,415（確率加重：Base55%+Bull20%+Bear25%）→ 現値比+38.2%"),
    ("FY2027/3 EPS推算", "¥95.1（NI¥270億÷2.84億株）"),
    ("ForwardPER（ガイダンスベース）", "33.6x（V1コンセンサスベース51.5xから劇的改善）"),
    ("配当", "FY2026/3 ¥42確定 / FY2027/3 ¥48予想（3期連続大幅増配）"),
    ("配当利回り（FY2027/3予想）", "1.50%（¥48÷¥3,195）"),
    ("EV/EBITDA", "18.5-20x（グローバル衛星ピア比プレミアム）"),
    ("PBR / ROE", "3.19x / ~8.2%（FY2026/3実績）→ FY2027/3で~9.5%へ改善見込み"),
    ("Beta", "0.176（低ボラティリティ＝防衛的特性）"),
    ("損切り水準", "¥2,800（25MA大幅割れ、▲12.4%）— V1から変更なし"),
    ("最大リスク", "PBR3.19x vs ROE9.5%乖離。防衛省コンステレーション受注遅延リスク"),
    ("投資テーマ", "GEO衛星スロット独占×ホルムズ地政学×日本防衛費2%GDP目標×トライサット"),
    ("IRコンタクト", "https://www.skyperfectjsat.space/jsat/ir/contact/"),
]

for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}：{value}")
    set_font(r, size=10.5)

doc.add_page_break()

# ── 本文パース ──
i = 0
table_buf = []
in_code   = False
code_buf  = []

while i < len(lines):
    raw      = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_font(run, size=9)
            run.font.name = "Courier New"
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    if stripped.startswith("|"):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    if stripped.startswith("### "):
        add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        add_heading(stripped[2:], level=1)
    elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        add_body(stripped.strip("*"), bold=True)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        content = stripped[2:]
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=10, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10)
    elif stripped.startswith("> "):
        add_body(stripped[2:], indent=True)
    elif stripped == "" or stripped == "---":
        doc.add_paragraph()
    else:
        add_body(stripped)

    i += 1

if table_buf:
    add_table_from_md(table_buf)

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
