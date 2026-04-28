"""Convert 4063 信越化学工業 DeepResearch MD → DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH = "outputs/fukai/4063_信越化学工業_STRONGBUY_DeepResearch_20260426.md"
DOCX_PATH = "outputs/fukai/4063_信越化学工業_STRONGBUY_DeepResearch_20260426.docx"

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    rPr.insert(0, rFonts)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: (0, 70, 127), 2: (31, 73, 125), 3: (54, 96, 146)}
    sizes = {1: 16, 2: 13, 3: 11.5}
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, (0, 0, 0)))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if level <= 2:
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "4472C4")
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=10, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table_from_md(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = t.cell(i, j)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            cell.text = clean
            runs = cell.paragraphs[0].runs
            if runs:
                set_font(runs[0], size=9, bold=(i == 0))
                if i == 0:
                    runs[0].font.color.rgb = RGBColor(255, 255, 255)
            if i == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F497D")
                tcPr.append(shd)
            elif i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DEEAF1")
                tcPr.append(shd)
    doc.add_paragraph()


# ── カバーページ ──────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("竹村トレードカンパニー アナリストチーム")
set_font(run, size=11, color=(46, 116, 181))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("4063 信越化学工業")
set_font(run, size=22, bold=True, color=(0, 70, 127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Deep Research — STRONG BUY +99")
set_font(run, size=14, bold=True, color=(46, 116, 181))

doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ("担当アナリスト", "深井"),
    ("分析日", "2026-04-26"),
    ("TP", "¥10,400（現値+53.7%）"),
    ("安全域", "+53.7%（現値¥6,768）"),
    ("優先度", "1（最高）"),
    ("次回決算", "2026年4月28日（月）15:30"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}：{value}")
    set_font(run, size=11)

doc.add_page_break()

# ── 本文パース ────────────────────────────────────────────
i = 0
table_buf = []
in_code = False
code_buf = []

while i < len(lines):
    raw = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped.startswith("```"):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_font(run, size=9)
            run.font.name = "Courier New"
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    if stripped.startswith("|"):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    if stripped.startswith("### "):
        add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        add_heading(stripped[2:], level=1)
    elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        add_body(stripped.strip("*"), bold=True)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        content = stripped[2:]
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=10, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10)
    elif stripped.startswith("> "):
        add_body(stripped[2:], indent=True)
    elif stripped == "" or stripped == "---":
        doc.add_paragraph()
    else:
        add_body(stripped)

    i += 1

if table_buf:
    add_table_from_md(table_buf)

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
