"""1850 南海辰村建設 Deep Research MD → DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH   = "outputs/fukai/1850_南海辰村建設_WATCH_DeepResearch_20260428.md"
DOCX_PATH = "outputs/fukai/1850_南海辰村建設_WATCH_DeepResearch_20260428.docx"

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_NAVY  = (0,   70,  127)
C_BLUE  = (31,  73,  125)
C_DEEP  = (54,  96,  146)
C_WHITE = (255, 255, 255)

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Yu Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "Yu Gothic")
    rPr.insert(0, rf)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    colors = {1: C_NAVY, 2: C_BLUE, 3: C_DEEP}
    sizes  = {1: 16, 2: 13, 3: 11.5}
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=colors.get(level, (0, 0, 0)))
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if level <= 2:
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "4472C4")
        pBdr.append(bot)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)

def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size=10, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=10, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)

def add_table_from_md(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = t.cell(i, j)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            cell.text = clean
            runs = cell.paragraphs[0].runs
            if runs:
                set_font(runs[0], size=9, bold=(i == 0))
                if i == 0:
                    runs[0].font.color.rgb = RGBColor(*C_WHITE)
            if i == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F497D")
                tcPr.append(shd)
            elif i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DEEAF1")
                tcPr.append(shd)
    doc.add_paragraph()

# ── カバーページ ──
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("オーキッドツリーキャピタル アナリストチーム")
set_font(r, size=11, color=C_BLUE)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("1850 南海辰村建設株式会社")
set_font(r, size=20, bold=True, color=C_NAVY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Deep Research — WATCH +68（土木官庁バックログ急増×無借金化×PBR0.68倍）")
set_font(r, size=14, bold=True, color=C_BLUE)

doc.add_paragraph()
doc.add_paragraph()

cover_items = [
    ("担当アナリスト", "深井"),
    ("分析日", "2026-04-28（FY2026/3本決算・増配・自社株買い発表当日）"),
    ("投資判断", "WATCH +68 / 打診20%可"),
    ("現在株価", "¥467（本日+4.5%、決算・還元発表日）"),
    ("Base TP（12ヶ月）", "¥650（+39.2%）— PBR0.95倍回復シナリオ"),
    ("Bull TP", "¥850（+82.0%）— 土木官庁バックログ消化×PBR1.25倍"),
    ("Bear TP", "¥350（▲25.1%）— 民間建築回復遅延×来期減益継続"),
    ("加重平均TP", "¥619（Base50%+Bull25%+Bear25%）→ 現値比+32.5%"),
    ("推奨アクション", "打診20%（¥460〜480）。本玉はFY2027/3 Q1（7月）土木工事進捗確認後"),
    ("FY2026/3 OP", "¥2,842M（+19.4% YoY・OPM 6.2%）— 減収増益構造"),
    ("FY2027/3 OP予想", "¥2,500M（▲12.0%）— 保守的ガイダンス・上振れ余地検討"),
    ("バックログ", "¥86,552M（+18.9%）— 土木官庁+536%・1.64年分確保"),
    ("有利子負債", "¥770M（前期¥5,980M比▲87%圧縮）→ 実質無借金"),
    ("現預金/株", "¥231（株価の49.5%）"),
    ("PBR / BPS", "0.68倍 / ¥683（解散価値以下）"),
    ("ROE（FY2026/3）", "11.3%（中堅ゼネコン比高水準）"),
    ("配当", "¥8（+33%増配）+ 自社株買い37億円・消却60万株を同日決議"),
    ("損切り水準", "¥380（52週安値¥290〜直近値¥467の中間）"),
    ("IRコンタクト", "https://www.nantatsu.co.jp / TEL: 06-6644-7805（経理部長 藤原）"),
]

for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}：{value}")
    set_font(r, size=10.5)

doc.add_page_break()

# ── 本文パース ──
i = 0
table_buf = []
in_code   = False
code_buf  = []

while i < len(lines):
    raw      = lines[i].rstrip("\n")
    stripped = raw.strip()

    if stripped.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            set_font(run, size=9)
            run.font.name = "Courier New"
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
        i += 1
        continue
    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    if stripped.startswith("|"):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        if table_buf:
            add_table_from_md(table_buf)
            table_buf = []

    if stripped.startswith("### "):
        add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        add_heading(stripped[2:], level=1)
    elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        add_body(stripped.strip("*"), bold=True)
    elif stripped.startswith("- ") or stripped.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        content = stripped[2:]
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=10, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10)
    elif stripped.startswith("> "):
        add_body(stripped[2:], indent=True)
    elif stripped == "" or stripped == "---":
        doc.add_paragraph()
    else:
        add_body(stripped)

    i += 1

if table_buf:
    add_table_from_md(table_buf)

doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
